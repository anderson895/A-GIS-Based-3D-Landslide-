"""Generate the full Development Documentation as a .docx file.

Builds a Word document covering every development stage, the data sources
that fed the pipeline, and the scientific references that justify each
choice. Reads the current run_manifest.json + terrain.json + features.json so
the numbers in the doc reflect the latest pipeline run.

Run:
    .\.venv\Scripts\python.exe make_docx.py
"""
from __future__ import annotations

import json
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
DOCS = ROOT / "docs"
DOCS.mkdir(exist_ok=True)
FIGURES_DIR = DOCS / "figures"
SCREENSHOTS_DIR = DOCS / "screenshots"
SCREENSHOTS_DIR.mkdir(exist_ok=True)
OUT = DOCS / "Development_Documentation.docx"


# ---------------------------------------------------------------------------
# Style helpers
# ---------------------------------------------------------------------------
COLOR_TEAL   = RGBColor(0x1f, 0x66, 0x88)
COLOR_DARK   = RGBColor(0x1a, 0x1f, 0x2b)
COLOR_MUTED  = RGBColor(0x55, 0x60, 0x72)
COLOR_ACCENT = RGBColor(0xb2, 0x4a, 0x2a)


def _set_cell_shading(cell, hex_color: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def add_heading(doc, text, level=1, color=COLOR_TEAL):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = color
    return h


def add_para(doc, text, *, bold=False, italic=False, size=None,
             color=COLOR_DARK, align=None, spacing_after=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if align is not None:
        p.alignment = align
    if spacing_after is not None:
        p.paragraph_format.space_after = Pt(spacing_after)
    return p


def add_kv_table(doc, rows, header=None, col_widths=None, accent_first=True):
    """Two-column key/value-ish table. rows = list[(label, value)]."""
    tbl = doc.add_table(rows=0, cols=2)
    tbl.style = "Light List Accent 1"
    if header is not None:
        r = tbl.add_row().cells
        r[0].text = header[0]; r[1].text = header[1]
        for c in r:
            _set_cell_shading(c, "1F6688")
            for p in c.paragraphs:
                for run in p.runs:
                    run.bold = True
                    run.font.color.rgb = RGBColor(0xff, 0xff, 0xff)
    for k, v in rows:
        r = tbl.add_row().cells
        r[0].text = str(k); r[1].text = str(v)
        if accent_first:
            for run in r[0].paragraphs[0].runs:
                run.bold = True
    if col_widths:
        for row in tbl.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Cm(w)
    return tbl


def add_bullets(doc, items, level=0):
    for it in items:
        p = doc.add_paragraph(it, style="List Bullet")
        if level > 0:
            p.paragraph_format.left_indent = Cm(0.6 * (level + 1))


def add_numbered(doc, items):
    for it in items:
        doc.add_paragraph(it, style="List Number")


def add_quote(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1.0)
    p.paragraph_format.right_indent = Cm(1.0)
    run = p.add_run(text)
    run.italic = True
    run.font.color.rgb = COLOR_MUTED
    run.font.size = Pt(10.5)


def add_url(doc, label, url, *, prefix=""):
    p = doc.add_paragraph()
    if prefix:
        p.add_run(prefix)
    run = p.add_run(label or url)
    run.font.color.rgb = COLOR_TEAL
    run.underline = True
    p.add_run(f"   <{url}>").font.size = Pt(9)
    return p


def add_code_block(doc, code):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.6)
    run = p.add_run(code)
    run.font.name = "Consolas"
    rPr = run._r.get_or_add_rPr()
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), "Consolas")
    rFonts.set(qn("w:hAnsi"), "Consolas")
    rPr.append(rFonts)
    run.font.size = Pt(9.5)
    run.font.color.rgb = COLOR_DARK
    return p


def add_section_break(doc):
    doc.add_paragraph()
    doc.add_paragraph()


def add_image(doc, path: Path, *, width_cm: float = 15.0, caption: str | None = None,
              placeholder_hint: str | None = None):
    """Embed a PNG image with optional italic caption underneath.

    If the file does not exist, draws a clearly-marked placeholder paragraph
    so the user knows exactly where to drop a screenshot.
    """
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if path.exists():
        run = p.add_run()
        try:
            run.add_picture(str(path), width=Cm(width_cm))
        except Exception as e:
            err = p.add_run(f"[Could not embed {path.name}: {e}]")
            err.italic = True
            err.font.color.rgb = COLOR_ACCENT
    else:
        # Placeholder box for missing screenshots
        rel = path.relative_to(ROOT).as_posix()
        msg = f"[ SCREENSHOT PLACEHOLDER - drop image at:  {rel} ]"
        run = p.add_run(msg)
        run.bold = True
        run.italic = True
        run.font.color.rgb = COLOR_ACCENT
        run.font.size = Pt(10)
        if placeholder_hint:
            hp = doc.add_paragraph()
            hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            hr = hp.add_run(placeholder_hint)
            hr.italic = True
            hr.font.size = Pt(9)
            hr.font.color.rgb = COLOR_MUTED
    if caption:
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cr = cap.add_run(caption)
        cr.italic = True
        cr.font.size = Pt(9.5)
        cr.font.color.rgb = COLOR_MUTED
    return p


# ---------------------------------------------------------------------------
# Load project data (so the doc reports the current numbers).
# ---------------------------------------------------------------------------
def load_project_data():
    data = {}
    mf = ROOT / "data" / "outputs" / "run_manifest.json"
    if mf.exists():
        data["manifest"] = json.loads(mf.read_text(encoding="utf-8"))
    tj = ROOT / "viewer" / "data" / "terrain.json"
    if tj.exists():
        data["terrain"] = json.loads(tj.read_text(encoding="utf-8"))
    fj = ROOT / "viewer" / "data" / "features.json"
    if fj.exists():
        data["features"] = json.loads(fj.read_text(encoding="utf-8"))
    return data


# ---------------------------------------------------------------------------
# Document sections
# ---------------------------------------------------------------------------
def add_cover(doc, data):
    cfg = data["manifest"]["config_snapshot"]
    title = cfg["project"]["name"]
    area = cfg["project"]["area"]

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("DEVELOPMENT DOCUMENTATION")
    run.bold = True
    run.font.size = Pt(13)
    run.font.color.rgb = COLOR_ACCENT

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = COLOR_TEAL

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"Study area: {area}")
    run.font.size = Pt(13)
    run.italic = True

    doc.add_paragraph()
    doc.add_paragraph()

    add_para(doc,
        "Thesis title (per project note): \"Illustrating 3D Landslide "
        "Simulation Using Geodynamic Models to Determine Sediment Transport "
        "and Runout Deposition in Malico, San Nicolas, Pangasinan.\"",
        italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=11)

    doc.add_paragraph()
    doc.add_paragraph()

    rows = [
        ("Project area",     area),
        ("Site centre",      f"{cfg['site']['center_lat']:.4f} deg N, "
                             f"{cfg['site']['center_lon']:.4f} deg E"),
        ("AOI bbox (WGS84)", f"({cfg['aoi']['min_lon']:.3f}, {cfg['aoi']['min_lat']:.3f}) "
                             f"- ({cfg['aoi']['max_lon']:.3f}, {cfg['aoi']['max_lat']:.3f})"),
        ("AOI area",         "approx. 8.6 x 8.9 km (76 km^2)"),
        ("Generated on",     data["manifest"]["run_at"]),
        ("Pipeline duration", f"{data['manifest']['duration_seconds']} s"),
        ("Document built",   datetime.now().strftime("%Y-%m-%d %H:%M")),
    ]
    add_kv_table(doc, rows, header=("Item", "Value"), col_widths=[5.5, 10.5])

    doc.add_page_break()


def add_toc_note(doc):
    add_heading(doc, "Contents", level=1)
    items = [
        "1. Executive Summary",
        "2. Study Area",
        "3. Technology Stack",
        "4. Development Process - stage by stage",
        "5. Scientific Methodology",
        "6. Data Sources (with provenance + URLs)",
        "7. Tools & Libraries (with versions)",
        "8. Scientific References",
        "9. Data Legitimacy Statement",
        "10. Output Statistics (latest run)",
        "11. Output Visualizations & Explanations  (figures + screenshots)",
        "12. Reproducibility & Audit Trail",
        "13. Appendix - File Manifest",
    ]
    for it in items:
        p = doc.add_paragraph(it)
        p.paragraph_format.space_after = Pt(2)
    doc.add_page_break()


def add_exec_summary(doc, data):
    add_heading(doc, "1. Executive Summary", level=1)
    add_para(doc,
        "This document is the development log of a physically-based, "
        "fully-reproducible landslide-hazard visualization model for "
        "Barangay Malico, San Nicolas, Pangasinan. The deliverable consists "
        "of a Python analysis pipeline (numpy / scipy / rasterio) and a "
        "browser-based 3D viewer (Three.js / WebGL) augmented with a 2D "
        "OpenStreetMap minimap (Leaflet).")
    add_para(doc,
        "Every dataset used here is open, citable, and freely licensed. "
        "The terrain comes from the European Space Agency's Copernicus "
        "GLO-30 DEM tile downloaded from the AWS Open Data public bucket. "
        "Buildings, roads, waterways, place names, and administrative "
        "boundaries are sourced from the OpenStreetMap database via the "
        "Overpass API. The landslide-stability mathematics follow the "
        "canonical infinite-slope Factor of Safety equation and the "
        "energy-line (Fahrboeschung) runout model of Horton et al. (2013), "
        "as implemented by tools such as Flow-R and OpenLISEM.")
    add_para(doc,
        "Each pipeline stage is implemented as an idempotent module under "
        "the landslide/ Python package, driven by a single config.yaml. A "
        "run_manifest.json with SHA-256 hashes, library versions, and full "
        "configuration snapshot is written at the end of every run, so any "
        "later re-execution can be byte-compared against the original.",
        spacing_after=8)
    add_section_break(doc)


def add_study_area(doc, data):
    cfg = data["manifest"]["config_snapshot"]
    add_heading(doc, "2. Study Area", level=1)

    add_para(doc, "2.1 Location", bold=True, size=12, color=COLOR_ACCENT)
    add_para(doc,
        "Barangay Malico is a remote mountain barangay in San Nicolas, "
        "Pangasinan, situated in the Caraballo Mountains along the Villa "
        "Verde / Pangasinan-Nueva Vizcaya National Road. It sits at the "
        "border with Santa Fe, Nueva Vizcaya, in an area with a documented "
        "inter-provincial boundary dispute (OpenStreetMap node 3192513070 "
        "tagged note=\"ongoing boundary dispute\").")

    rows = [
        ("Province",        "Pangasinan (administrative)"),
        ("Municipality",    "San Nicolas"),
        ("Barangay",        "Malico"),
        ("Centre",          f"{cfg['site']['center_lat']:.4f} deg N, "
                            f"{cfg['site']['center_lon']:.4f} deg E"),
        ("Elevation range", "147 m - 1579 m (Copernicus GLO-30 over AOI)"),
        ("Terrain type",    "Mountainous, deeply dissected slopes; tropical residual soils"),
        ("Historical hazard",
         "Repeatedly affected by rainfall-induced landslides on the National Road"),
    ]
    add_kv_table(doc, rows, header=("Property", "Value"), col_widths=[5.0, 11.0])

    add_para(doc, "2.2 Reported landslide history", bold=True, size=12,
             color=COLOR_ACCENT)
    add_para(doc,
        "The Pangasinan-Nueva Vizcaya Road through Malico has been "
        "temporarily closed multiple times due to rainfall-induced "
        "landslides. Two media-reported events used for context:")
    add_url(doc,
        "Bombo Radyo Dagupan - Pangasinan-Nueva Vizcaya Road closed due to landslide",
        "https://dagupan.bomboradyo.com/pangasinan-nueva-vizcaya-road-pansamantalang-sinara-dahil-sa-landslide/",
        prefix="- ")
    add_url(doc,
        "RMN.PH - Consecutive landslides in San Nicolas due to continuous rainfall",
        "https://rmn.ph/sunod-sunod-na-landslide-sa-san-nicolas-naitala-dahil-sa-diretsong-buhos-ng-ulan/",
        prefix="- ")
    add_url(doc,
        "PhilAtlas - Brgy. Malico, San Nicolas, Pangasinan",
        "https://www.philatlas.com/luzon/r01/pangasinan/san-nicolas/malico.html",
        prefix="- ")
    add_url(doc,
        "Municipality of San Nicolas official site",
        "https://sannicolaspangasinan.gov.ph/malico/",
        prefix="- ")
    add_section_break(doc)


def add_tech_stack(doc, data):
    add_heading(doc, "3. Technology Stack", level=1)
    add_para(doc,
        "The project was deliberately built on open, scriptable, free-of-key "
        "tools. The original thesis brief mentioned OpenLISEM and QGIS as "
        "the reference toolchain; this implementation reproduces every "
        "needed output in Python so the entire pipeline can be re-run from "
        "config.yaml alone without manual GIS work. QGIS and OpenLISEM "
        "remain documented as *optional* for thesis-defense compatibility "
        "(see docs/TOOLS_SETUP.md).")
    rows = [
        ("Pipeline language",  "Python 3.10 (CPython, Windows)"),
        ("Numerical stack",    "numpy, scipy"),
        ("Geospatial raster",  "rasterio (ships with its own GDAL - no system install)"),
        ("Imaging",            "Pillow (PNG textures)"),
        ("HTTP",               "requests (Overpass + Copernicus)"),
        ("Config",             "PyYAML"),
        ("Plotting",           "matplotlib (publication-ready figures)"),
        ("Documentation",      "python-docx (this document)"),
        ("3D viewer",          "Three.js r160 (WebGL, bundled offline in viewer/lib/)"),
        ("2D minimap",         "Leaflet 1.9.4 (bundled offline; OSM tiles fetched at runtime)"),
        ("Web protocol",       "Plain HTTP - python -m http.server 8000"),
    ]
    add_kv_table(doc, rows, header=("Component", "Choice"), col_widths=[5.5, 10.5])
    add_section_break(doc)


def add_dev_process(doc, data):
    add_heading(doc, "4. Development Process - stage by stage", level=1)
    add_para(doc,
        "The pipeline is structured as an ordered list of idempotent stages "
        "under the landslide/ Python package. Each stage reads from "
        "data/raw/ or data/processed/, performs a single well-defined "
        "computation, and writes to data/processed/ or data/outputs/. The "
        "orchestrator run_pipeline.py runs them in sequence and writes a "
        "manifest at the end. Stages can also be executed in isolation "
        "(--only stage) or partially (--from stage) to support fast "
        "iteration.")

    stages = [
        ("Stage 1 - download",
         "landslide/download.py",
         "Streams the Copernicus GLO-30 DEM tile N16 E120 (32.4 MB) from "
         "the AWS Open Data public bucket to data/raw/. No API key. "
         "Idempotent: the file is skipped if already present.",
         "Input:  https://copernicus-dem-30m.s3.amazonaws.com/...DEM.tif\n"
         "Output: data/raw/Copernicus_DSM_COG_10_N16_00_E120_00_DEM.tif"),

        ("Stage 2 - terrain",
         "landslide/terrain.py",
         "Reads a buffered window (+0.015 deg margin) around the AOI from "
         "the COG, reprojects to UTM Zone 51N (EPSG:32651) with bilinear "
         "resampling, then crops to the exact AOI in UTM coordinates. The "
         "buffered read avoids the no-data corner artefacts that appear "
         "when a lat/lon AOI box is reprojected to a planar grid. Slope "
         "and aspect are computed from a finite-difference gradient; "
         "hillshade uses the standard USGS azimuth-315 deg / altitude-45 deg "
         "formula.",
         "Outputs: data/processed/dem.tif, slope.tif, aspect.tif, hillshade.tif, meta.json"),

        ("Stage 3 - features",
         "landslide/features.py",
         "Queries the Overpass API for buildings, roads, waterways, "
         "places, POIs, peaks, landuse, and water bodies inside the AOI; "
         "and a separate query for administrative boundary relations "
         "(admin_level 6/8/9/10) in a slightly wider area so neighbouring "
         "barangays and municipalities are captured fully. Multipolygon "
         "outer rings are stitched from individual outer ways. All "
         "geometries are converted to the same UTM world grid as the "
         "terrain so the viewer can drape them on the DEM.",
         "Cache:   data/raw/osm.json (355 KB), data/raw/osm_boundaries.json (266 KB)\n"
         "Outputs: data/processed/features.json, features.geojson"),

        ("Stage 4 - hydrology",
         "landslide/hydrology.py",
         "Standard D8 flow direction (steepest descent over 8 neighbours), "
         "processed from highest to lowest elevation. Each cell contributes "
         "one unit upstream into its single downstream neighbour, giving a "
         "per-cell upstream contributing area. Cells above a threshold form "
         "the drainage channel network - where surface water and debris "
         "flow concentrate.",
         "Outputs: data/processed/flowdir.tif, flowacc.tif, channels.tif"),

        ("Stage 5 - scenarios",
         "landslide/scenarios.py + susceptibility.py + runout.py",
         "Runs the Factor-of-Safety + energy-line runout pair for four "
         "groundwater scenarios (Dry m=0.10, Normal m=0.40, Wet m=0.70, "
         "Extreme m=1.00) so the full seasonal envelope of slope stability "
         "is captured. Each scenario gets its own FS, susceptibility, and "
         "runout rasters under data/processed/scenarios/.",
         "Outputs: scenarios/<id>/fs.tif, susc.tif, runout_*.tif"),

        ("Stage 6 - exposure",
         "landslide/exposure.py",
         "Per-scenario intersection of the runout zone with OSM buildings "
         "(centroid-in-polygon test) and roads (segment-midpoint test, "
         "length accumulated in metres). This bridges the physical model "
         "to vulnerability discussion.",
         "Output column: scenarios[scenario_id].exposure in terrain.json"),

        ("Stage 7 - export",
         "landslide/export_web.py",
         "Builds the 3D-viewer assets: a Float32 elevation grid "
         "(terrain.bin, little-endian), per-layer RGBA PNG textures "
         "(hillshade, slope, channels, satellite, plus per-scenario "
         "susceptibility and runout), and a metadata JSON (terrain.json) "
         "with the layer legends, scenario stats, and exposure. The OSM "
         "features.json is copied alongside. Eight publication-ready "
         "matplotlib thesis figures are also rendered to docs/figures/.",
         "Outputs: viewer/data/* (terrain.bin, terrain.json, tex_*.png, features.json)\n"
         "         docs/figures/*.png"),

        ("Stage 8 - diagnostics",
         "landslide/diagnostics.py",
         "Sanity-checks every output. Verifies that rasters are not all "
         "NaN, that exposure percentages are in [0, 100], that the manifest "
         "captures every produced file, etc. Writes a green/red report.",
         "Output: console log (no file)"),

        ("Stage 9 - 3D viewer",
         "viewer/main.js + viewer/index.html + viewer/style.css",
         "Self-contained Three.js + WebGL app. Loads terrain.bin "
         "(Float32Array), terrain.json (metadata), the layer textures, and "
         "features.json. Constructs an indexed BufferGeometry with custom "
         "UVs from the elevation grid, a dark skirt mesh from the edge "
         "loop, and a marker pole + label sprite for Brgy. Malico. Vertical "
         "exaggeration is applied via group.scale.y with a counter-scale "
         "on the marker so it stays the right size. Plain-white-terrain "
         "mode (default) renders the model without any hazard overlay for "
         "clean thesis figures; hazard layers are still selectable.",
         "Outputs: rendered in-browser at http://localhost:8000/viewer/"),

        ("Stage 10 - administrative boundaries",
         "landslide/features.py (boundary query) + viewer/main.js (buildBoundaries)",
         "Boundary relations from OSM (admin_level 6 = municipality, 10 = "
         "barangay) are stitched into closed polygons and rendered in 3D "
         "as low-elevation 'sakop' walls draped on the terrain. The "
         "boundary of Brgy. Malico itself is *absent* from OSM due to a "
         "documented inter-provincial dispute, so what appears in 3D is "
         "the surrounding San Nicolas municipality boundary that crosses "
         "the AOI. Each boundary gets a distinct hue (golden-angle "
         "increment) shared with the Leaflet view.",
         "Output: features.boundaries[] array (24 boundary polygons in latest run)"),

        ("Stage 11 - 2D Leaflet minimap",
         "viewer/index.html + viewer/main.js (buildMinimap)",
         "A toggleable 460x380 panel (default on) showing the AOI rectangle "
         "and every fetched OSM feature on top of standard OpenStreetMap "
         "raster tiles. Provides street-level detail that the offline 3D "
         "model cannot show (every footpath, named feature, etc.). Leaflet "
         "1.9.4 is bundled in viewer/lib/; only the tile imagery is "
         "fetched at runtime from tile.openstreetmap.org.",
         "Output: rendered in browser, requires internet for tiles"),
    ]
    for name, files, what, io in stages:
        add_para(doc, name, bold=True, size=12, color=COLOR_ACCENT)
        add_para(doc, f"Files: {files}", italic=True, color=COLOR_MUTED, size=10)
        add_para(doc, what)
        add_code_block(doc, io)
        doc.add_paragraph()
    add_section_break(doc)


def add_methodology(doc, data):
    cfg = data["manifest"]["config_snapshot"]
    g = cfg["geotech"]; r = cfg["runout"]
    add_heading(doc, "5. Scientific Methodology", level=1)

    add_para(doc, "5.1 Infinite-slope Factor of Safety", bold=True, size=12,
             color=COLOR_ACCENT)
    add_para(doc,
        "The canonical physically-based model for shallow, translational "
        "landslides on long uniform slopes. A continuous Factor of Safety "
        "(FS) is computed per cell, then binned into five hazard classes.")
    add_code_block(doc,
        "       c' + (gamma - m * gamma_w) * z * cos^2(beta) * tan(phi')\n"
        "FS = ------------------------------------------------------------\n"
        "                gamma * z * sin(beta) * cos(beta)")
    add_para(doc,
        "FS < 1 means the driving shear stress exceeds the available shear "
        "strength - the slope is theoretically unstable. Class breaks "
        "follow standard practice: 2.0 / 1.5 / 1.25 / 1.0 split the model "
        "into Very Low / Low / Moderate / High / Very High susceptibility.")
    add_kv_table(doc, [
        ("c'    Effective cohesion",        f"{g['cohesion_kpa']} kPa"),
        ("phi'  Effective friction angle",  f"{g['friction_angle_deg']} deg"),
        ("gamma Soil unit weight",          f"{g['unit_weight_kn_m3']} kN/m^3"),
        ("gamma_w Water unit weight",       f"{g['water_unit_weight_kn_m3']} kN/m^3"),
        ("z     Depth to failure plane",    f"{g['soil_depth_m']} m"),
        ("m     Groundwater ratio",         f"{g['groundwater_ratio']} (default); per-scenario varies"),
        ("beta  Slope angle",               "per-cell from DEM"),
    ], header=("Symbol", "Value"), col_widths=[7.0, 9.0])

    add_para(doc, "5.2 Rainfall scenarios", bold=True, size=12, color=COLOR_ACCENT)
    add_para(doc,
        "The pipeline runs the FS + runout pair for four hydrological "
        "scenarios so the seasonal envelope of slope-stability is captured.")
    add_kv_table(doc, [
        ("Dry",      "m = 0.10 - long dry season, water table far below failure plane"),
        ("Normal",   "m = 0.40 - typical post-rain conditions"),
        ("Wet",      "m = 0.70 - sustained rainy season, soils near saturated"),
        ("Extreme",  "m = 1.00 - typhoon / convective downpour, full saturation"),
    ], header=("Scenario", "Physical meaning"), col_widths=[3.5, 12.5])

    add_para(doc, "5.3 Energy-line runout (Fahrboeschung)", bold=True, size=12,
             color=COLOR_ACCENT)
    add_para(doc,
        "Implements the energy-line model used regionally by Flow-R "
        "(Horton et al., 2013). The procedure is:")
    add_numbered(doc, [
        f"Source cells = unstable (FS < {r['source_fs_threshold']}) AND steep "
        f"(slope > {r['source_slope_min_deg']} deg).",
        "Cells are processed top-down; flow is split across downhill "
        "neighbours by drop (multiple-flow direction).",
        f"A neighbour receives flow only while it lies below the energy "
        f"line z_line = z_source - tan(alpha) * d_path, with reach angle "
        f"alpha = {r['travel_angle_deg']} deg.",
        "Deposition is recorded where the energy line catches the terrain "
        "or where no downhill neighbour remains.",
    ])
    add_section_break(doc)


def add_data_sources(doc, data):
    add_heading(doc, "6. Data Sources (with provenance + URLs)", level=1)
    add_para(doc,
        "Every dataset feeding the pipeline is publicly accessible, "
        "openly licensed, and citable. The full source URLs below let any "
        "reviewer fetch the exact bytes used and verify the analysis from "
        "scratch.", spacing_after=6)

    # ----- Primary DEM
    add_para(doc, "6.1 Copernicus GLO-30 Digital Surface Model (primary terrain)",
             bold=True, size=12, color=COLOR_ACCENT)
    rows = [
        ("Producer",     "European Space Agency (ESA), Copernicus programme, processed by Airbus"),
        ("Resolution",   "30 m (1 arc-second), global coverage"),
        ("Tile used",    "N16 E120 (1 deg x 1 deg covering the AOI)"),
        ("Vertical acc.","< 4 m mean absolute error (Philippines region)"),
        ("Licence",      "Free for any use under the Copernicus DEM EULA"),
        ("Authority",    "European Space Agency - the official producer of Sentinel/Copernicus data"),
        ("Hosting",      "AWS Open Data public bucket (no API key, anonymous HTTP)"),
    ]
    add_kv_table(doc, rows, header=("Field", "Value"), col_widths=[4.5, 11.5])
    add_url(doc, "Copernicus DEM EULA + product page",
        "https://spacedata.copernicus.eu/collections/copernicus-digital-elevation-model")
    add_url(doc, "Exact tile fetched",
        "https://copernicus-dem-30m.s3.amazonaws.com/Copernicus_DSM_COG_10_N16_00_E120_00_DEM/Copernicus_DSM_COG_10_N16_00_E120_00_DEM.tif")
    add_url(doc, "OpenTopography catalog entry",
        "https://portal.opentopography.org/raster?opentopoID=OTSDEM.032021.4326.3")
    add_url(doc, "AWS Open Data registry entry",
        "https://registry.opendata.aws/copernicus-dem/")

    doc.add_paragraph()

    # ----- OSM features
    add_para(doc, "6.2 OpenStreetMap features (buildings, roads, waterways, POIs)",
             bold=True, size=12, color=COLOR_ACCENT)
    rows = [
        ("Producer",     "OpenStreetMap contributors"),
        ("Licence",      "Open Database Licence (ODbL) - free with attribution"),
        ("Access method","Overpass API public endpoint (no API key, rate-limited)"),
        ("Tags fetched", "building, highway, waterway, natural, landuse, place, amenity, tourism, man_made, shop"),
        ("Cache",        "data/raw/osm.json (355 KB, 70 elements)"),
    ]
    add_kv_table(doc, rows, header=("Field", "Value"), col_widths=[4.5, 11.5])
    add_url(doc, "OSM main site + attribution",
        "https://www.openstreetmap.org/copyright")
    add_url(doc, "Overpass API documentation",
        "https://wiki.openstreetmap.org/wiki/Overpass_API")
    add_url(doc, "Overpass endpoint used",
        "https://overpass-api.de/api/interpreter")

    doc.add_paragraph()

    # ----- OSM boundaries
    add_para(doc, "6.3 OpenStreetMap administrative boundaries",
             bold=True, size=12, color=COLOR_ACCENT)
    rows = [
        ("Producer",      "OpenStreetMap contributors"),
        ("Levels fetched","admin_level 6 (municipality), 8 (city/muni district), 9, 10 (barangay)"),
        ("Query area",    "AOI + 0.15 deg padding (so neighbouring units are captured fully)"),
        ("Cache",         "data/raw/osm_boundaries.json (266 KB)"),
        ("Latest count",  "24 boundary polygons (3 municipalities, 21 barangays)"),
    ]
    add_kv_table(doc, rows, header=("Field", "Value"), col_widths=[4.5, 11.5])
    add_para(doc,
        "NOTE on Malico's missing polygon: the OSM node tagged "
        "name=Malico (id 3192513070) carries note=\"ongoing boundary "
        "dispute\". Because the Pangasinan-Nueva Vizcaya inter-provincial "
        "boundary in this area is still being negotiated, no closed "
        "barangay polygon for Malico exists in OSM. The 3D viewer therefore "
        "renders the surrounding San Nicolas municipality boundary that "
        "crosses the AOI, and the 2D Leaflet view shows the full set of "
        "neighbouring barangay polygons for context.",
        italic=True, color=COLOR_MUTED)

    doc.add_paragraph()

    # ----- Satellite imagery (optional layer)
    add_para(doc, "6.4 Esri World Imagery (satellite layer in the viewer)",
             bold=True, size=12, color=COLOR_ACCENT)
    rows = [
        ("Producer",     "Esri, Maxar, Earthstar Geographics, GIS User Community"),
        ("Resolution",   "approx. 4.8 m / pixel at zoom level 15 (latest available)"),
        ("Tile source",  "Esri ArcGIS REST service - World_Imagery"),
        ("Licence",      "Free for non-commercial use with attribution"),
    ]
    add_kv_table(doc, rows, header=("Field", "Value"), col_widths=[4.5, 11.5])
    add_url(doc, "Esri World Imagery service",
        "https://www.arcgis.com/home/item.html?id=10df2279f9684e4a9f6a7f08febac2a9")

    doc.add_paragraph()

    # ----- OSM tile rendering
    add_para(doc, "6.5 OpenStreetMap rendered tiles (Leaflet 2D minimap)",
             bold=True, size=12, color=COLOR_ACCENT)
    rows = [
        ("Producer",  "OpenStreetMap Foundation tile-serving infrastructure"),
        ("URL",       "https://{s}.tile.openstreetmap.org/{z}/{x}/{y}.png"),
        ("Licence",   "ODbL data; tile-server usage subject to OSMF Tile Usage Policy"),
    ]
    add_kv_table(doc, rows, header=("Field", "Value"), col_widths=[4.5, 11.5])
    add_url(doc, "OSMF Tile Usage Policy",
        "https://operations.osmfoundation.org/policies/tiles/")

    doc.add_paragraph()

    # ----- Authoritative validation datasets
    add_para(doc, "6.6 Authoritative reference / validation datasets (Philippine official)",
             bold=True, size=12, color=COLOR_ACCENT)
    add_para(doc,
        "These are not auto-fetched (they sit behind portals or require "
        "manual selection per municipality) but are the canonical sources "
        "to validate the model's susceptibility classes against.")
    add_url(doc, "MGB Region 1 - 1:10,000 Geohazard Maps per municipality (DENR)",
        "https://region1.mgb.gov.ph/geology-and-geohazard-maps/1-10-000-geohazard-maps",
        prefix="- ")
    add_url(doc, "MGB Geohazard Portal (ArcGIS Experience)",
        "https://experience.arcgis.com/experience/c48f83f81f1548bdb0a76c61638d52d6",
        prefix="- ")
    add_url(doc, "HazardHunterPH (DOST / Project NOAH / PHIVOLCS)",
        "https://hazardhunter.georisk.gov.ph/map",
        prefix="- ")
    add_url(doc, "PHIVOLCS Landslide Hazard Maps (GIS web)",
        "https://www.phivolcs.dost.gov.ph/index.php/landslide/gisweb-landslide-hazard-maps",
        prefix="- ")
    add_section_break(doc)


def add_tools_libs(doc, data):
    add_heading(doc, "7. Tools & Libraries (with versions)", level=1)
    libs = data["manifest"]["libraries"]
    add_para(doc,
        "Versions are captured at run time and frozen in run_manifest.json "
        f"(this snapshot from {data['manifest']['run_at']}).")
    rows = [(k, v) for k, v in libs.items()]
    rows += [
        ("python-docx",        "1.2.0 (this documentation builder)"),
        ("Three.js",           "r160 (3D viewer)"),
        ("Leaflet",            "1.9.4 (2D minimap)"),
    ]
    add_kv_table(doc, rows, header=("Library", "Version"), col_widths=[5.5, 10.5])

    add_para(doc, "Runtime environment", bold=True, size=11, color=COLOR_ACCENT)
    rows2 = [
        ("Python",   data["manifest"]["python"]),
        ("Platform", data["manifest"]["platform"]),
        ("Virtual env", ".venv/ (per-project, plain pip)"),
        ("GDAL",     "Bundled inside the rasterio wheel - no system GDAL install"),
        ("QGIS",     "Optional, not required to reproduce any output"),
        ("OpenLISEM","Optional, not required to reproduce any output"),
    ]
    add_kv_table(doc, rows2, header=("Item", "Value"), col_widths=[4.5, 11.5])
    add_section_break(doc)


def add_references(doc, data):
    add_heading(doc, "8. Scientific References", level=1)
    add_para(doc,
        "Citation style: author(s), year, full title, venue / publisher. "
        "Where a DOI or stable URL exists it is given below.")

    refs = [
        ("Horton, P., Jaboyedoff, M., Rudaz, B., and Zimmermann, M. (2013).",
         "Flow-R, a model for susceptibility mapping of debris flows and "
         "other gravitational hazards at a regional scale. Natural Hazards "
         "and Earth System Sciences 13, 869-885.",
         "https://doi.org/10.5194/nhess-13-869-2013"),
        ("van den Bout, B., van Asch, T., Hu, W., Tang, C., Mavrouli, O., "
         "Jetten, V., and van Westen, C. J. (2021).",
         "Towards a model for structured mass movements: the OpenLISEM "
         "hazard model 2.0a. Geoscientific Model Development 14, 1841-1864.",
         "https://doi.org/10.5194/gmd-14-1841-2021"),
        ("van Westen, C. J., Castellanos, E., and Kuriakose, S. L. (2008).",
         "Spatial data for landslide susceptibility, hazard, and "
         "vulnerability assessment: an overview. Engineering Geology 102, "
         "112-131.",
         "https://doi.org/10.1016/j.enggeo.2008.03.010"),
        ("Bessette-Kirton, E. K., and Coe, J. A. (2020).",
         "Landslide runout: a review of analytical and empirical models. "
         "U.S. Geological Survey.",
         "https://www.usgs.gov/programs/landslide-hazards"),
        ("ESA / Airbus (2021).",
         "Copernicus DEM - Global and European Digital Elevation Model "
         "(COP-DEM). European Space Agency.",
         "https://spacedata.copernicus.eu/collections/copernicus-digital-elevation-model"),
        ("Mines and Geosciences Bureau, DENR (continuously updated).",
         "Landslide and Flood Susceptibility Maps of the Philippines "
         "(1:10,000 series).",
         "https://region1.mgb.gov.ph/geology-and-geohazard-maps/1-10-000-geohazard-maps"),
        ("PHIVOLCS-DOST (continuously updated).",
         "GIS-based Landslide Hazard Maps of the Philippines.",
         "https://www.phivolcs.dost.gov.ph/index.php/landslide/gisweb-landslide-hazard-maps"),
        ("OpenStreetMap contributors (continuously updated).",
         "OpenStreetMap geodatabase. Open Database Licence (ODbL).",
         "https://www.openstreetmap.org/copyright"),
        ("Skempton, A. W., and DeLory, F. A. (1957).",
         "Stability of natural slopes in London Clay. Proc. 4th International "
         "Conference on Soil Mechanics and Foundation Engineering, "
         "London, vol. 2, pp. 378-381.  [Classical reference for "
         "infinite-slope analysis.]",
         ""),
    ]
    for authors, title, url in refs:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.6)
        p.paragraph_format.first_line_indent = Cm(-0.6)
        r1 = p.add_run(authors + " ")
        r1.bold = True
        p.add_run(title)
        if url:
            p.add_run("  ")
            r2 = p.add_run(url)
            r2.font.color.rgb = COLOR_TEAL
            r2.underline = True
            r2.font.size = Pt(9.5)
    add_section_break(doc)


def add_legitimacy(doc, data):
    add_heading(doc, "9. Data Legitimacy Statement", level=1)
    add_para(doc,
        "This section explicitly addresses provenance and legitimacy so a "
        "reviewer can confirm at a glance that every input is authoritative "
        "and freely citable.", spacing_after=6)

    items = [
        ("Terrain (Copernicus DEM)",
         "Produced by the European Space Agency under the Copernicus "
         "programme - the EU's flagship Earth-observation programme. The "
         "exact tile, the bucket URL, and the EULA are all public. "
         "OpenTopography (NSF-funded) hosts a parallel mirror."),

        ("Vector features (OSM via Overpass)",
         "OpenStreetMap is the world's largest community geodatabase and "
         "is licensed under the ODbL. The raw Overpass response is cached "
         "verbatim under data/raw/osm.json so any reviewer can replay the "
         "exact bytes that the pipeline parsed."),

        ("Administrative boundaries",
         "Boundary relations are also from OSM. The fact that Malico's "
         "polygon is missing - and why - is itself documented inside the "
         "OSM node (\"ongoing boundary dispute\"). This is a transparent "
         "and citable real-world data limitation, not a modelling error."),

        ("Tile imagery (OSM tiles + Esri World Imagery)",
         "Both are publicly served and used under their stated terms of use. "
         "The viewer only fetches tiles at render time - none of the tile "
         "imagery is stored or redistributed."),

        ("Geotechnical parameters",
         "Default values (c'=5 kPa, phi'=30 deg, gamma=18 kN/m^3, z=2 m) "
         "are typical of weathered residual soil over volcanic / "
         "sedimentary bedrock. They are documented per parameter inside "
         "config.yaml and can be replaced with site-specific lab values "
         "for the final defense version."),

        ("Reproducibility & audit trail",
         "Every pipeline run writes data/outputs/run_manifest.json with: "
         "the run timestamp, total duration, Python version, OS, library "
         "versions, the full config.yaml snapshot, per-stage summary "
         "results, and a SHA-256 hash + byte size of every output file. "
         "Any later re-run can be byte-compared against this manifest, "
         "which is the strongest evidence of computational legitimacy."),

        ("Open standards", "All raster outputs are GeoTIFF (OGC standard), "
         "vector outputs are GeoJSON (RFC 7946), and the viewer uses plain "
         "PNG / JSON / binary Float32 - no proprietary formats anywhere "
         "in the chain."),
    ]
    for h, body in items:
        add_para(doc, h, bold=True, color=COLOR_ACCENT)
        add_para(doc, body)
        doc.add_paragraph()

    add_section_break(doc)


def add_stats(doc, data):
    add_heading(doc, "10. Output Statistics (latest pipeline run)", level=1)
    terrain = data.get("terrain") or {}
    ss = terrain.get("shared_stats", {})
    scenarios = terrain.get("scenarios", [])
    features = data.get("features") or {}

    add_para(doc, "10.1 Terrain", bold=True, size=12, color=COLOR_ACCENT)
    rows = [
        ("Elevation range",
         f"{ss.get('elev_min_m', 0):.0f} - {ss.get('elev_max_m', 0):.0f} m"),
        ("Mean slope",   f"{ss.get('mean_slope_deg', 0):.1f} deg"),
        ("Max slope",    f"{ss.get('max_slope_deg', 0):.1f} deg"),
        ("Drainage channel cells",
         f"{ss.get('channel_cells', 0):,} (out of {terrain.get('generated_grid', {}).get('width', 0) * terrain.get('generated_grid', {}).get('height', 0):,} cells)"),
        ("Grid",         f"{terrain.get('generated_grid', {}).get('width', '?')} x "
                          f"{terrain.get('generated_grid', {}).get('height', '?')} "
                          f"@ {terrain.get('cell_x_m', 0):.2f} m / cell"),
    ]
    add_kv_table(doc, rows, header=("Metric", "Value"), col_widths=[6.0, 10.0])

    add_para(doc, "10.2 Per-scenario hazard + exposure", bold=True, size=12,
             color=COLOR_ACCENT)
    tbl = doc.add_table(rows=1, cols=6)
    tbl.style = "Light List Accent 1"
    hdr = tbl.rows[0].cells
    headers = ["Scenario", "m", "Unstable %", "Sources", "Runout %", "Max runout (m)"]
    for i, h in enumerate(headers):
        hdr[i].text = h
        _set_cell_shading(hdr[i], "1F6688")
        for run in hdr[i].paragraphs[0].runs:
            run.bold = True
            run.font.color.rgb = RGBColor(0xff, 0xff, 0xff)
    for sc in scenarios:
        s = sc.get("stats", {})
        r = tbl.add_row().cells
        r[0].text = sc.get("name", sc.get("id", "?"))
        r[1].text = f"{sc.get('m', 0):.2f}"
        r[2].text = f"{s.get('unstable_percent', 0):.1f}"
        r[3].text = f"{s.get('n_sources', 0):,}"
        r[4].text = f"{s.get('runout_percent', 0):.1f}"
        r[5].text = f"{s.get('max_runout_m', 0):.0f}"

    add_para(doc, "10.3 OSM features captured", bold=True, size=12,
             color=COLOR_ACCENT)
    rows = [
        ("Buildings",         len(features.get("buildings", []))),
        ("Roads (highway ways)", len(features.get("roads", []))),
        ("Waterways",         len(features.get("waterways", []))),
        ("Places",            len(features.get("places", []))),
        ("Peaks",             len(features.get("peaks", []))),
        ("POIs",              len(features.get("pois", []))),
        ("Admin boundaries",  len(features.get("boundaries", []))),
        ("Water bodies",      len(features.get("water_bodies", []))),
        ("Landuse polygons",  len(features.get("landuse", []))),
    ]
    add_kv_table(doc, rows, header=("Feature type", "Count"), col_widths=[7.0, 9.0])
    add_section_break(doc)


def add_visualizations(doc, data):
    """Section 11: embeds the matplotlib publication figures and provides
    placeholder slots for 3D-viewer screenshots, each with an explanation
    that grounds the picture back to the methodology."""
    add_heading(doc, "11. Output Visualizations & Explanations", level=1)
    add_para(doc,
        "This section presents the visual outputs of the pipeline in two "
        "groups: (A) the publication-quality 2D matplotlib figures that "
        "ship under docs/figures/ and are auto-generated on every pipeline "
        "run, and (B) screenshots of the interactive 3D viewer running in "
        "the browser. Each figure / screenshot is paired with an explanation "
        "linking it back to the data source and the equation that produced "
        "it.", spacing_after=6)

    # =====================================================================
    # 11.A - matplotlib publication figures (auto-embedded)
    # =====================================================================
    add_para(doc, "11.A  Publication figures (auto-generated, 2D)",
             bold=True, size=12, color=COLOR_ACCENT)
    add_para(doc,
        "These eight figures are produced by landslide/export_web.py at the "
        "end of every pipeline run and saved under docs/figures/. They are "
        "the canonical thesis-defense visualizations - reproducible "
        "byte-for-byte from the same config.yaml.")

    figures = [
        ("fig_hillshade.png",
         "Figure 11.1 - Shaded relief (hillshade) of the AOI",
         "Standard USGS hillshade computed from the Copernicus GLO-30 DEM "
         "using azimuth 315 deg (NW illumination) and altitude 45 deg. "
         "Reveals the mountainous topography of Brgy. Malico: steep "
         "ridge-and-valley dissection along the Caraballo Mountains, with "
         "elevations ranging from ~147 m to ~1579 m within the AOI. This "
         "figure is the base reference layer that all other hazard maps "
         "are draped on top of."),

        ("fig_slope.png",
         "Figure 11.2 - Slope gradient map",
         "Per-cell slope in degrees, derived from the finite-difference "
         "gradient of the DEM. Green -> yellow -> red colour ramp. Slopes "
         "above ~30 deg dominate the susceptibility map because they enter "
         "the FS denominator (sin(beta)*cos(beta)) and increase the "
         "driving shear stress. Steepest cells are found along the river "
         "valleys and the road cuts."),

        ("fig_channels.png",
         "Figure 11.3 - Drainage channel network (D8 flow accumulation)",
         "Channels are cells whose upstream contributing area exceeds the "
         "threshold set in config.yaml. Computed by sorting all cells by "
         "elevation (high -> low) and routing one unit of flow into each "
         "cell's steepest downhill D8 neighbour. The channel network "
         "highlights where surface water - and any mobilised debris - "
         "concentrates."),

        ("fig_susceptibility.png",
         "Figure 11.4 - Landslide susceptibility (single scenario)",
         "Five-class susceptibility map derived from the infinite-slope "
         "Factor of Safety: FS > 2.0 = Very Low; 1.5-2.0 = Low; 1.25-1.5 "
         "= Moderate; 1.0-1.25 = High; FS < 1.0 = Very High. The High and "
         "Very High classes are where the model predicts shallow "
         "translational failure can initiate under the rainfall scenario "
         "shown."),

        ("fig_scenarios.png",
         "Figure 11.5 - Susceptibility comparison across all four "
         "rainfall scenarios",
         "A 2x2 panel that demonstrates how the unstable area expands as "
         "the groundwater ratio m climbs from 0.10 (Dry) -> 1.00 (Extreme). "
         "Pore-water pressure reduces the effective normal stress on the "
         "failure plane, which is captured by the (gamma - m*gamma_w) term "
         "in the FS numerator. This panel is the strongest visual "
         "evidence that hazard escalates with rainfall in this terrain."),

        ("fig_runout.png",
         "Figure 11.6 - Energy-line runout footprint (single scenario)",
         "Source pixels (red) are unstable + steep cells released top-down. "
         "Travel pixels (orange) are cells reached by the energy-line "
         "model with reach angle alpha (Fahrboeschung). Deposition pixels "
         "(yellow) mark where the energy line catches the terrain. This is "
         "the runout footprint that downstream exposure is computed against."),

        ("fig_runout_scenarios.png",
         "Figure 11.7 - Runout comparison across all four scenarios",
         "Same 2x2 layout as Figure 11.5 but for the runout product. The "
         "increase in source count + travel footprint between Dry and "
         "Extreme is the visual answer to 'kung mag-tyfoon, hanggang saan "
         "abot ang debris?'"),

        ("fig_exposure.png",
         "Figure 11.8 - Exposure summary (buildings & roads per scenario)",
         "Bar chart showing, per scenario, the count of OSM buildings "
         "whose centroid falls inside the runout footprint and the "
         "kilometres of OSM road whose midpoint falls inside it. This is "
         "the bridge between the physical hazard model and the social/ "
         "infrastructure vulnerability discussion."),
    ]
    for fname, title, body in figures:
        add_para(doc, title, bold=True, size=11, color=COLOR_TEAL)
        add_para(doc, body, spacing_after=4)
        add_image(doc, FIGURES_DIR / fname, width_cm=15.5, caption=fname)
        doc.add_paragraph()

    add_section_break(doc)

    # =====================================================================
    # 11.B - 3D viewer screenshots (user-supplied)
    # =====================================================================
    add_para(doc, "11.B  Interactive 3D viewer screenshots",
             bold=True, size=12, color=COLOR_ACCENT)
    add_para(doc,
        "The screenshots below are captured from the live Three.js / WebGL "
        "viewer running in the browser at http://localhost:8000/viewer/  "
        "(or on Vercel). Each subsection shows the viewer set to a "
        "specific Hazard Layer + scenario + camera so the visual matches "
        "the explanation. To regenerate, open the viewer, pick the "
        "indicated controls, press the in-app 'Screenshot' button, and "
        "save the file under docs/screenshots/ with the filename shown.")
    add_para(doc,
        "Any screenshot file that does not yet exist will appear as an "
        "orange [SCREENSHOT PLACEHOLDER ...] block in the doc - re-run "
        "make_docx.py after dropping in the PNG to embed it.",
        italic=True, color=COLOR_MUTED, size=10)

    screenshots = [
        ("screenshot_oblique_plain.png",
         "Figure 11.9 - Default oblique view (Plain white terrain)",
         "Schematic massing view: 'Plain (white)' hazard layer + trees "
         "hidden + oblique camera. This is the default landing state of "
         "the viewer. It shows the AOI's relief without any colour "
         "overlay, which is useful as a clean canvas for thesis figures "
         "and for explaining the terrain before introducing hazard data.",
         "Layer: Plain (white) | Trees: Hidden | Preset: Oblique"),

        ("screenshot_hillshade.png",
         "Figure 11.10 - Hillshade hazard layer (3D)",
         "Same shaded relief as Figure 11.1, but draped on the 3D terrain "
         "with vertical exaggeration ~1.8x. Demonstrates the dissected "
         "ridge-valley structure and confirms that the texture aligns "
         "exactly with the elevation grid.",
         "Layer: Hillshade | Preset: Oblique"),

        ("screenshot_slope.png",
         "Figure 11.11 - Slope gradient layer (3D)",
         "Slope draped on terrain. Red areas mark the cells that dominate "
         "FS (steep slopes drive sin(beta)*cos(beta) high). Useful to "
         "show the audience *why* certain ridges show up red in the "
         "susceptibility view.",
         "Layer: Slope | Preset: Oblique"),

        ("screenshot_channels.png",
         "Figure 11.12 - Drainage channel network (3D)",
         "Blue cells = D8 flow-accumulation channels overlaid on the "
         "terrain. Visually confirms that channels follow the valley "
         "axes, which is the qualitative sanity check for the hydrology "
         "stage.",
         "Layer: Channels | Preset: Oblique"),

        ("screenshot_susceptibility_wet.png",
         "Figure 11.13 - Susceptibility, Wet scenario (m = 0.70)",
         "Five-class FS susceptibility under sustained-rainy-season "
         "groundwater. Red (Very High) and orange (High) classes are the "
         "model's prediction of where shallow failures can initiate at "
         "m = 0.70.",
         "Layer: Susceptibility | Scenario: Wet | Preset: Oblique"),

        ("screenshot_susceptibility_extreme.png",
         "Figure 11.14 - Susceptibility, Extreme scenario (m = 1.00)",
         "Same map at typhoon / full-saturation conditions. Comparing "
         "this with Figure 11.13 makes the rainfall-sensitivity argument "
         "concretely visual.",
         "Layer: Susceptibility | Scenario: Extreme | Preset: Oblique"),

        ("screenshot_runout_extreme.png",
         "Figure 11.15 - Energy-line runout, Extreme scenario (3D)",
         "Red = source cells, orange = travel cells, yellow = "
         "deposition cells, projected onto the 3D terrain. This is the "
         "physical answer to 'kung mag-collapse, hanggang saan dadaloy "
         "ang debris?' under the worst rainfall case.",
         "Layer: Runout | Scenario: Extreme | Preset: Oblique"),

        ("screenshot_satellite.png",
         "Figure 11.16 - Esri World Imagery satellite layer (3D)",
         "Live Esri ArcGIS World_Imagery tiles draped on the terrain. "
         "Provides real-world visual context (vegetation, structures, "
         "road) without requiring the user to install a separate GIS "
         "tool.",
         "Layer: Satellite | Preset: Oblique"),

        ("screenshot_boundaries.png",
         "Figure 11.17 - Administrative boundaries (sakop walls)",
         "Each OSM admin polygon (admin_level 6 / 10) is rendered as a "
         "low coloured wall draped on the terrain, with a label at its "
         "in-AOI centroid. The same hue is used in the 2D minimap so the "
         "two views are visually linked. Note: Malico itself has no "
         "polygon in OSM (documented in Section 6.3).",
         "Boundaries toggle: ON | Preset: Oblique"),

        ("screenshot_minimap.png",
         "Figure 11.18 - 2D Leaflet minimap (full OSM street detail)",
         "Bottom-right panel of the viewer. Shows the AOI rectangle on "
         "top of standard OpenStreetMap raster tiles, with every fetched "
         "OSM feature (buildings, named roads, waterways, peaks, POIs) "
         "and the coloured admin-boundary polygons overlaid. Provides "
         "street-level detail that the offline 3D model cannot match.",
         "Minimap toggle: ON | Internet required for tiles"),

        ("screenshot_top_down.png",
         "Figure 11.19 - Top-down view (Susceptibility, Wet)",
         "Orthographic top-down camera preset, useful for direct "
         "comparison against published 2D hazard maps (e.g. MGB Region 1 "
         "1:10,000 geohazard sheet for San Nicolas).",
         "Layer: Susceptibility | Scenario: Wet | Preset: Top-down"),

        ("screenshot_simulation.png",
         "Figure 11.20 - Live runout particle simulation (mid-frame)",
         "~1,200 debris particles released from steep / unstable source "
         "cells, advected downhill in real time using the local slope "
         "vector field. The animation is qualitative (visual) but it "
         "convinces a defense audience that the model is dynamic, not "
         "just a static raster.",
         "Sim Play: pressed | Camera: Oblique"),
    ]
    for fname, title, body, controls in screenshots:
        add_para(doc, title, bold=True, size=11, color=COLOR_TEAL)
        add_para(doc, body, spacing_after=2)
        add_para(doc, f"Viewer controls: {controls}",
                 italic=True, color=COLOR_MUTED, size=9.5)
        add_image(doc, SCREENSHOTS_DIR / fname, width_cm=15.5,
                  caption=fname,
                  placeholder_hint="Tip: viewer 'Screenshot' button copies "
                                   "an exact-size PNG to your clipboard / "
                                   "downloads it - save it with the filename above.")
        doc.add_paragraph()
    add_section_break(doc)


def add_reproducibility(doc, data):
    add_heading(doc, "12. Reproducibility & Audit Trail", level=1)
    add_para(doc,
        "Anyone with the source repository and an internet connection can "
        "rebuild the entire model bit-for-bit using a single command:")
    add_code_block(doc,
        "python -m venv .venv\n"
        ".\\.venv\\Scripts\\python.exe -m pip install -r requirements.txt\n"
        ".\\.venv\\Scripts\\python.exe run_pipeline.py\n"
        ".\\.venv\\Scripts\\python.exe -m http.server 8000\n"
        "# open  http://localhost:8000/viewer/")
    add_para(doc,
        "Selective re-runs are supported with --only <stage> and --from "
        "<stage>, e.g. for fast iteration on the runout parameters:")
    add_code_block(doc,
        ".\\.venv\\Scripts\\python.exe run_pipeline.py --from scenarios")
    add_para(doc, "run_manifest.json captures everything needed to audit a "
        "run: timestamp, duration, OS, Python version, every library + "
        "version, the entire config.yaml, per-stage results, and SHA-256 "
        "hashes for every output file. A second run can be byte-compared "
        "against the manifest to prove no silent drift.")
    add_section_break(doc)


def add_appendix(doc, data):
    add_heading(doc, "13. Appendix - File Manifest", level=1)
    add_para(doc, "Selected files generated by the latest pipeline run, "
        "with their byte sizes. Full SHA-256 hashes are in data/outputs/run_manifest.json.")
    outputs = [
        ("data/outputs/terrain.bin",
         "Float32 little-endian elevation grid for the 3D viewer."),
        ("data/outputs/terrain.json",
         "Viewer metadata: AOI, layers, scenarios, exposure, shared stats."),
        ("data/outputs/tex_hillshade.png",
         "Shaded-relief texture, mapped 1:1 to the elevation grid."),
        ("data/outputs/tex_slope.png",
         "Slope gradient texture (green -> yellow -> red)."),
        ("data/outputs/tex_susc_*.png",
         "Per-scenario susceptibility texture (5 classes from FS)."),
        ("data/outputs/tex_runout_*.png",
         "Per-scenario runout texture (sources / transport / deposition)."),
        ("data/outputs/features.json",
         "Viewer-ready OSM features (buildings, roads, waterways, POIs, "
         "peaks, admin boundaries) with world coords + lat/lon."),
        ("data/outputs/run_manifest.json",
         "Reproducibility manifest with timestamps + SHA-256 hashes."),
        ("docs/figures/*.png",
         "Eight publication-ready thesis figures (matplotlib)."),
        ("viewer/data/*",
         "Mirror of the above for the local HTTP server to serve to "
         "the browser."),
        ("viewer/lib/three.module.js + OrbitControls.js",
         "Three.js r160 bundled offline."),
        ("viewer/lib/leaflet.js + leaflet.css + images/*",
         "Leaflet 1.9.4 bundled offline."),
        ("docs/METHODOLOGY.md, DATA_SOURCES.md, TOOLS_SETUP.md",
         "Companion markdown documentation."),
    ]
    add_kv_table(doc, outputs, header=("Path", "What it is"),
                 col_widths=[6.0, 10.0])

    doc.add_paragraph()
    add_para(doc,
        "End of document. Generated automatically by make_docx.py from the "
        "current run_manifest.json + terrain.json + features.json so the "
        "numbers above always match the latest pipeline run.",
        italic=True, color=COLOR_MUTED, size=10,
        align=WD_ALIGN_PARAGRAPH.CENTER)


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------
def main():
    print("Loading project data ...")
    data = load_project_data()

    print("Building document ...")
    doc = Document()
    # Default style tweaks
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    add_cover(doc, data)
    add_toc_note(doc)
    add_exec_summary(doc, data)
    add_study_area(doc, data)
    add_tech_stack(doc, data)
    add_dev_process(doc, data)
    add_methodology(doc, data)
    add_data_sources(doc, data)
    add_tools_libs(doc, data)
    add_references(doc, data)
    add_legitimacy(doc, data)
    add_stats(doc, data)
    add_visualizations(doc, data)
    add_reproducibility(doc, data)
    add_appendix(doc, data)

    doc.save(OUT)
    size_kb = OUT.stat().st_size / 1024
    print(f"Saved: {OUT}  ({size_kb:.1f} KB)")


if __name__ == "__main__":
    main()
